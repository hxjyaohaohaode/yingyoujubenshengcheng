import os
import logging
from dataclasses import dataclass

from sqlalchemy.ext.asyncio import AsyncSession

from core.narrative.script_parser import parse_script_content, build_narrative_memory_from_script
from core.narrative.style_analyzer import analyze_style, get_style_guide

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}
MAX_FILE_SIZE = 20 * 1024 * 1024


@dataclass
class ParsedDocument:
    title: str
    text: str
    page_count: int = 1
    source_path: str = ""
    file_type: str = ""


class FileUploadHandler:
    def validate_file(self, filename: str, file_size: int) -> tuple[bool, str]:
        ext = os.path.splitext(filename)[1].lower()
        if ext not in ALLOWED_EXTENSIONS:
            return False, f"不支持的文件格式 '.{ext}'。支持: PDF、Word、TXT、Markdown"
        if file_size > MAX_FILE_SIZE:
            return False, f"文件过大 ({file_size / (1024 * 1024):.1f}MB)。单个文件上限20MB"
        if file_size == 0:
            return False, "文件为空"
        return True, ""

    async def parse_file(self, file_content: bytes, filename: str) -> ParsedDocument:
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            return await self._parse_pdf(file_content, filename)
        elif ext == ".docx":
            return await self._parse_docx(file_content, filename)
        elif ext in (".txt", ".md", ".markdown"):
            return await self._parse_text(file_content, filename)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    async def _parse_pdf(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            import io
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
            text = "\n\n".join(text_parts)
            return ParsedDocument(title=os.path.splitext(filename)[0], text=text,
                                  page_count=len(reader.pages), source_path=filename, file_type="pdf")
        except ImportError:
            try:
                import io
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    text_parts = [page.extract_text() for page in pdf.pages if page.extract_text()]
                text = "\n\n".join(text_parts)
                return ParsedDocument(title=os.path.splitext(filename)[0], text=text,
                                      page_count=len(text_parts), source_path=filename, file_type="pdf")
            except ImportError:
                raise ImportError("需要安装 PyPDF2 或 pdfplumber 来解析PDF文件")

    async def _parse_docx(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(content))
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n".join(text_parts)
            return ParsedDocument(title=os.path.splitext(filename)[0], text=text,
                                  page_count=len(text_parts) // 40 + 1, source_path=filename, file_type="docx")
        except ImportError:
            raise ImportError("需要安装 python-docx 来解析Word文件")

    async def _parse_text(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = content.decode("gbk")
            except UnicodeDecodeError:
                text = content.decode("latin-1")
        lines = [l for l in text.split("\n") if l.strip()]
        return ParsedDocument(title=os.path.splitext(filename)[0], text="\n".join(lines),
                              page_count=len(lines) // 50 + 1, source_path=filename,
                              file_type=os.path.splitext(filename)[1].lstrip("."))


async def parse_uploaded_script(db: AsyncSession, project_id: str, file_content: str, filename: str = "") -> dict:
    parsed = parse_script_content(file_content, filename)
    style = analyze_style(file_content)
    style_guide = get_style_guide(style)

    memory_result = await build_narrative_memory_from_script(db, project_id, parsed)

    return {
        "title": filename,
        "total_words": parsed.total_words,
        "chapter_count": len(parsed.chapters),
        "character_count": len(parsed.characters),
        "characters": [c["name"] for c in parsed.characters],
        "characters_detail": parsed.characters,
        "chapters": [{"index": ch["index"], "title": ch["title"], "word_count": ch["word_count"]} for ch in parsed.chapters],
        "style": {
            "avg_sentence_length": style.avg_sentence_length,
            "dialogue_ratio": style.dialogue_ratio,
            "narrative_pov": style.narrative_pov,
            "tone_keywords": style.tone_keywords,
            "summary": style.summary,
            "style_guide": style_guide,
        },
        "memory_initialized": memory_result,
    }